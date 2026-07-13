"""
image_generator.py — AI-powered image generation for business reports.

Generates context-aware images for gap analysis and narrative reports using
RouteLLM gateway (Midjourney, DALL-E 3, GPT-4V image edit).

Usage:
    from rag_pipeline.image_generator import generate_report_images
    images = generate_report_images(
        narrative_text="Gap analysis for Process Network...",
        intent="process",
        report_type="gap_analysis"
    )
    # Returns: [{"url": "...", "prompt": "...", "model": "midjourney"}, ...]
"""

from __future__ import annotations

import logging
import os
import re
import time
from typing import Optional

import requests

from .langfuse_tracer import get_tracer

logger = logging.getLogger(__name__)

# RouteLLM gateway config
_ROUTELLM_BASE_URL = os.environ.get("HF_BASE_URL")
_ROUTELLM_API_KEY = os.environ.get("HF_TOKEN")
_IMAGE_GENERATION_BUDGET = int(os.environ.get("IMAGE_GENERATION_BUDGET", "10"))
_IMAGE_TIMEOUT = int(os.environ.get("IMAGE_GENERATION_TIMEOUT", "90"))

# Model availability (mapped to RouteLLM modalities)
_IMAGE_MODELS = {
    "process": "midjourney",      # Process diagrams benefit from Midjourney's visual style
    "people": "midjourney",         # People/org charts work well with DALL-E
    "technology": "midjourney",   # Tech stack visualizations → Midjourney
    "general": "midjourney",
}


def _extract_key_concepts(narrative_text: str, intent: str, max_concepts: int = 3) -> list[str]:
    """Extract key concepts from narrative text for image prompt engineering.

    Args:
        narrative_text: Full narrative/analysis text
        intent: "process", "people", "technology", or "general"
        max_concepts: Max concepts to extract

    Returns:
        List of concept phrases for image generation
    """
    # Remove markdown/formatting
    clean = re.sub(r"[*_#\-\n]+", " ", narrative_text)

    # Intent-specific extraction patterns
    patterns = {
        "process": [
            r"(?:bottleneck|flow|cycle time|lead time|critical path)[^.]*",
            r"(?:workflow|process|improvement)[^.]*",
            r"(?:optimization|efficiency|throughput)[^.]*",
        ],
        "people": [
            r"(?:team|role|collaboration|ownership|stakeholder)[^.]*",
            r"(?:organizational|structure|communication)[^.]*",
            r"(?:resource|capacity|engagement)[^.]*",
        ],
        "technology": [
            r"(?:platform|tool|integration|automation|pipeline)[^.]*",
            r"(?:infrastructure|deployment|ci/cd|devops)[^.]*",
            r"(?:modernization|migration|stack)[^.]*",
        ],
    }

    concepts = []
    for pattern in patterns.get(intent, []):
        matches = re.findall(pattern, clean, re.IGNORECASE)
        for m in matches:
            m = m.strip()
            if len(m) > 10 and m not in concepts:
                concepts.append(m)
                if len(concepts) >= max_concepts:
                    break
        if len(concepts) >= max_concepts:
            break

    return concepts[:max_concepts]


def _build_image_prompt(
    narrative_text: str,
    intent: str,
    report_type: str,
    concept: str,
) -> str:
    """Build a structured image prompt from concept and context.

    Args:
        narrative_text: Full narrative for context
        intent: "process", "people", "technology"
        report_type: "gap_analysis", "assessment", "roadmap"
        concept: Key concept to visualize

    Returns:
        Prompt string optimized for image generation
    """
    prefixes = {
        "process": "Professional business process diagram showing",
        "people": "Organizational chart and team structure illustrating",
        "technology": "Technology stack and infrastructure diagram depicting",
        "general": "Business concept visualization of",
    }

    prefix = prefixes.get(intent, prefixes["general"])

    # Build context-aware prompt
    if report_type == "gap_analysis":
        prompt = (
            f"{prefix} gaps and improvements in: {concept}. "
            f"Style: modern, professional, business report, clean design, "
            f"muted colors, minimal text labels. High-resolution business graphic."
        )
    elif report_type == "assessment":
        prompt = (
            f"{prefix} current state and maturity of: {concept}. "
            f"Style: professional assessment visualization, business-appropriate, "
            f"clean layout, assessment metrics visible. Corporate design."
        )
    elif report_type == "roadmap":
        prompt = (
            f"{prefix} roadmap and transformation timeline for: {concept}. "
            f"Style: strategic roadmap visualization, modern business, "
            f"timeline elements visible, forward-looking design."
        )
    else:
        prompt = (
            f"{prefix}: {concept}. "
            f"Style: professional, business-appropriate, clean, corporate design."
        )

    return prompt


def generate_report_images(
    narrative_text: str,
    intent: str = "general",
    report_type: str = "assessment",
    max_images: int | None = None,
) -> list[dict]:
    """Generate AI images for narrative reports using RouteLLM gateway.

    Args:
        narrative_text: Full narrative/analysis text
        intent: "process", "people", "technology", or "general"
        report_type: "gap_analysis", "assessment", "roadmap", etc.
        max_images: Max images to generate (defaults to _IMAGE_GENERATION_BUDGET)

    Returns:
        List of dicts: {"url": str, "prompt": str, "model": str, "concept": str}
        Empty list if generation disabled or fails.

    Gracefully handles:
        - Missing API key
        - Budget exhaustion
        - Request failures
        - Image generation failures
    """
    if not _ROUTELLM_API_KEY:
        logger.warning("IMAGE_GENERATION: No API key configured. Skipping image generation.")
        return []

    if not narrative_text or len(narrative_text.strip()) < 50:
        logger.debug("IMAGE_GENERATION: Narrative too short. Skipping images.")
        return []

    max_images = max_images or _IMAGE_GENERATION_BUDGET
    if max_images <= 0:
        logger.debug("IMAGE_GENERATION: Budget is 0. Skipping images.")
        return []

    # Extract key concepts to visualize
    concepts = _extract_key_concepts(narrative_text, intent, max_concepts=max_images)
    if not concepts:
        logger.warning("IMAGE_GENERATION: Could not extract concepts from narrative.")
        return []

    model = _IMAGE_MODELS.get(intent, _IMAGE_MODELS["general"])
    results = []

    tracer = get_tracer()
    endpoint = f"{_ROUTELLM_BASE_URL}/chat/completions"

    for i, concept in enumerate(concepts[:max_images], 1):
        prompt = _build_image_prompt(narrative_text, intent, report_type, concept)
        logger.info(
            "IMAGE_GENERATION: Generating image %d/%d via %s (concept: %s)",
            i, len(concepts), model, concept[:50],
        )

        with tracer.trace(
            "routellm_image_generation",
            input={
                "endpoint": endpoint,
                "model": model,
                "concept_preview": concept[:80],
                "intent": intent,
                "report_type": report_type,
                "image_index": i,
            },
        ) as span:
            try:
                _t0 = time.monotonic()
                response = requests.post(
                    endpoint,
                    headers={
                        "Authorization": f"Bearer {_ROUTELLM_API_KEY}",
                        "Content-Type": "application/json",
                    },
                    json={
                        "model": model,
                        "messages": [{"role": "user", "content": prompt}],
                        "modalities": ["image"],
                        "temperature": 0.7,
                        "max_tokens": 1024,
                    },
                    timeout=_IMAGE_TIMEOUT,
                )
                _latency_ms = round((time.monotonic() - _t0) * 1000, 1)

                if response.status_code != 200:
                    logger.warning(
                        "IMAGE_GENERATION: HTTP %d from %s: %s",
                        response.status_code, model, response.text[:200],
                    )
                    span.set_output({
                        "status_code": response.status_code,
                        "latency_ms": _latency_ms,
                        "images_returned": 0,
                        "success": False,
                    })
                    span.set_error(f"HTTP {response.status_code}: {response.text[:200]}")
                    continue

                data = response.json()
                if data.get("error"):
                    _err_msg = data["error"].get("message", str(data["error"]))[:200]
                    logger.warning("IMAGE_GENERATION: %s error: %s", model, _err_msg)
                    span.set_output({
                        "status_code": response.status_code,
                        "latency_ms": _latency_ms,
                        "images_returned": 0,
                        "success": False,
                        "api_error": _err_msg,
                    })
                    span.set_error(_err_msg)
                    continue

                # Extract image URLs from response
                _images_this_call = 0
                for choice in data.get("choices", []):
                    msg = choice.get("message", {})
                    content = msg.get("content", [])
                    if isinstance(content, str):
                        content = [{"type": "text", "text": content}]

                    for item in content:
                        if item.get("type") == "image_url":
                            url = item.get("image_url", {}).get("url")
                            if url:
                                results.append({
                                    "url": url,
                                    "prompt": prompt,
                                    "model": model,
                                    "concept": concept,
                                })
                                _images_this_call += 1
                                logger.info(
                                    "IMAGE_GENERATION: Image generated via %s (%s)",
                                    model, concept[:40],
                                )

                span.set_output({
                    "status_code": response.status_code,
                    "latency_ms": _latency_ms,
                    "images_returned": _images_this_call,
                    "success": True,
                })

                if i < len(concepts):
                    time.sleep(2)

            except requests.Timeout:
                logger.warning("IMAGE_GENERATION: Timeout on image %d/%d", i, len(concepts))
                span.set_error(f"Timeout after {_IMAGE_TIMEOUT}s")
            except Exception as e:
                logger.warning("IMAGE_GENERATION: Failed on concept '%s': %s", concept[:40], e)
                span.set_error(str(e))

    logger.info("IMAGE_GENERATION: Generated %d images for %s report", len(results), report_type)
    return results


def embed_images_in_document(
    doc_path: str,
    images: list[dict],
    placement: str = "inline",
) -> bool:
    """Embed generated images into a Word document.

    Args:
        doc_path: Path to .docx file
        images: List of image dicts from generate_report_images()
        placement: "inline" (after sections) or "appendix" (at end)

    Returns:
        True if successful, False on error
    """
    if not images:
        return True

    try:
        from docx import Document
        from docx.shared import Inches
    except ImportError:
        logger.warning("python-docx not installed. Skipping image embedding.")
        return False

    try:
        doc = Document(doc_path)

        if placement == "appendix":
            # Add appendix heading
            doc.add_heading("Generated Visualizations", level=2)

        tracer = get_tracer()
        for img_data in images:
            try:
                url = img_data["url"]
                concept = img_data["concept"]
                model = img_data.get("model", "")

                # Download image from URL
                with tracer.trace(
                    "image_download",
                    input={"url_preview": url[:80], "concept": concept[:60], "model": model},
                ) as dl_span:
                    _t0 = time.monotonic()
                    img_response = requests.get(url, timeout=30)
                    _latency_ms = round((time.monotonic() - _t0) * 1000, 1)
                    dl_span.set_output({
                        "status_code": img_response.status_code,
                        "latency_ms": _latency_ms,
                        "bytes": len(img_response.content) if img_response.status_code == 200 else 0,
                        "success": img_response.status_code == 200,
                    })
                    if img_response.status_code != 200:
                        dl_span.set_error(f"HTTP {img_response.status_code}")

                if img_response.status_code != 200:
                    logger.warning("Failed to download image from %s", url[:80])
                    continue

                # Create temp file
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                    tmp.write(img_response.content)
                    tmp_path = tmp.name

                # Add to document
                doc.add_paragraph(f"Visualization: {concept} ({model})", style="Caption")
                doc.add_picture(tmp_path, width=Inches(5.5))
                doc.add_paragraph()  # Spacing

                # Cleanup
                os.unlink(tmp_path)

                logger.info("Embedded image for concept: %s", concept[:40])

            except Exception as e:
                logger.warning("Failed to embed image for '%s': %s", concept[:40], e)

        doc.save(doc_path)
        logger.info("Document saved with %d images", len(images))
        return True

    except Exception as e:
        logger.error("Failed to embed images in document: %s", e)
        return False
